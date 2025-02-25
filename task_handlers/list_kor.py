import logging
import win32print
import win32api
import os
from docx import Document
from conn.conn import bot
from conn.conn import get_db_connection
from task_handlers.delete_pair import delete_expired_message_pairs
from task_handlers.delete_pair import delete_all_message_pairs
from task_handlers.delete_pair import save_message_pair
from task_handlers.menu import main_menu
from telebot import types


logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")


TASKS_DOCX_PATH = "tasks.docx"

TEMPLATE_DOCX_PATH = "templates.docx"

def create_default_template():
    """Создает стандартный шаблон, если он отсутствует."""
    doc = Document()
    doc.add_heading('Список задач подрядчика', 0)

    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер задачи'
    hdr_cells[1].text = 'Наименование задачи'
    hdr_cells[2].text = 'Местоположение'
    hdr_cells[3].text = 'Комментарий'
    hdr_cells[4].text = 'Отчет о работе подрядчика'

    #
    doc.save(TEMPLATE_DOCX_PATH)

@bot.message_handler(func=lambda message: message.text == "📑Королев")
def contractor_task_list_handler(message):
    delete_expired_message_pairs(message.chat.id)
    delete_all_message_pairs(message.chat.id)
    save_message_pair(message.chat.id, user_message_id=message.message_id)

    sent_message = bot.send_message(message.chat.id, "Загрузка списка задач подрядчика...")
    save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)

    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, task_name, location, comment FROM contractor_tasks")
            tasks = cursor.fetchall()

            if tasks:
                
                task_list_message = "📑 Список задач подрядчика:\n\n"
                for task in tasks:
                    task_id, task_name, location, comment = task
                    task_list_message += (
                        f"🟡 Номер задачи: {task_id}\n"
                        f"   Наименование: {task_name}\n"
                        f"   Местоположение: {location}\n"
                        f"   Комментарий: {comment}\n\n"
                    )

                
                sent_message = bot.send_message(message.chat.id, task_list_message)
                save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)

                # Создаем или загружаем шаблон
                if not os.path.exists(TEMPLATE_DOCX_PATH):
                    create_default_template()
                    logging.info("Создан стандартный шаблон.")

                
                doc = Document(TEMPLATE_DOCX_PATH)

                
                table = None
                for tbl in doc.tables:
                    if len(tbl.columns) == 5:  
                        table = tbl
                        break

                
                if not table:
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Номер задачи'
                    hdr_cells[1].text = 'Наименование задачи'
                    hdr_cells[2].text = 'Местоположение'
                    hdr_cells[3].text = 'Комментарий'
                    hdr_cells[4].text = 'Отчет о работе подрядчика'

                
                for task in tasks:
                    task_id, task_name, location, comment = task
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(task_id)
                    row_cells[1].text = task_name
                    row_cells[2].text = location
                    row_cells[3].text = comment
                    row_cells[4].text = ""  

                
                doc.save(TASKS_DOCX_PATH)

                
                sent_message = bot.send_message(message.chat.id, "Список задач сохранен в Word-документ.")
                save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)

                
                sent_message = bot.send_message(
                    message.chat.id,
                    "Для печати списка задач подрядчика нажмите на команду /print"
                )
                save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)
            else:
                sent_message = bot.send_message(message.chat.id, "❌ Задач для подрядчика не найдено.")
                save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)
                main_menu(message)
    except Exception as e:
        logging.error(f"Ошибка при получении задач подрядчика: {str(e)}", exc_info=True)
        sent_message = bot.send_message(message.chat.id, "🚫 Не удалось загрузить задачи. Попробуйте снова.")
        save_message_pair(message.chat.id, bot_message_id=sent_message.message_id)

@bot.message_handler(commands=['print'])
def print_tasks_handler(message):
    try:
        
        if not os.path.exists(TASKS_DOCX_PATH):
            bot.send_message(message.chat.id, "❌ Файл с задачами не найден. Сначала загрузите задачи.")
            return

        
        printer_name = win32print.GetDefaultPrinter()
        win32api.ShellExecute(0, "print", TASKS_DOCX_PATH, f'"{printer_name}"', ".", 0)

        bot.send_message(message.chat.id, "🖨️ Список задач отправлен на печать.")
    except Exception as e:
        logging.error(f"Ошибка при печати задач: {str(e)}", exc_info=True)
        bot.send_message(message.chat.id, "🚫 Не удалось отправить задачи на печать.")

def main_menu_markup():
    """Формирует кнопки для главного меню."""
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add("➕ Добавить", "📋 Списки", "⚙️ Сервис")
    return markup

if __name__ == "__main__":
    logging.info("Бот запущен.")
    bot.polling(none_stop=True, timeout=30)
